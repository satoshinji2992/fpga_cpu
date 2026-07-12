#!/usr/bin/env python3
"""Update the course-design DOCX with the verified R16 results."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
REPORT = next(ROOT.glob("*.docx"))


def replace_paragraph(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_cell(cell, text):
    replace_paragraph(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        replace_paragraph(paragraph, "")


doc = Document(REPORT)

updates = {
    17: "时间：2026  年  7  月  12  日",
    23: "本课程设计面向“基于 FPGA 开发板的处理器设计”题目要求，完成了一套可在 TEC-PLUS 核心板上独立运行的 RV32 小型计算机系统。系统以五级流水线处理器为主核，集成 16 KiB 片上指令 ROM、4 KiB 数据 RAM、2 路组相联 I-Cache、64 MiB SDRAM、UART、LED、按键、机器态中断和性能计数器，并由统一板端固件提供交互式 shell、8×8 手写数字识别、浮点表达式计算器、Pong 和 SDRAM Paint 应用。",
    24: "系统采用 Icarus Verilog 自动回归与 FPGA 真板验证相结合的方法。R16 已通过 17 项自动化回归；复位后 CPU 执行十项开机自动测试并把结果独立保存到 m0～m9。50 MHz 真板验证中开机自检、数字识别、浮点计算器、Pong、Paint、SDRAM 与中断功能均已通过。",
    32: "本项目完成基础层次和进阶层次全部核心内容，并在拓展层次实现五级流水线冒险处理、动态分支预测、2 路组相联 Cache、RV32M 乘除法、Custom Float32、位操作扩展、两层浮点 MLP、表达式计算器、性能计数器和 PPA 分析框架。",
    36: "构建 16 KiB 指令 ROM、4 KiB 片内数据 RAM、直接映射/2 路组相联 I-Cache，以及两片 HY57V2562 组成的 64 MiB、32-bit SDRAM 子系统；",
    38: "实现统一板端固件及交互式串口 shell，支持 status、m0～m9、irq、sdram、perf、led、cnn、calc、pong、paint 等命令；",
    40: "实现 8×8 手写数字识别：CPU 接收 64 个二值像素，使用 Custom Float32 权重执行 64→8→10 两层 MLP（ReLU）并返回预测数字；精确板端模型测试集准确率为 82.84%，十个演示原型全部正确；",
    43: "建立 17 项自动化回归测试，完成 ISE 综合、布局布线、bitstream 下载及 50 MHz 真板串口联调。",
    47: "顶层模块 top.v 完成 CPU、I-Cache、片内 RAM、SDRAM 控制器与 MMIO 外设互联。TEC-PLUS 的 50 MHz 晶振经 BUFG 直接驱动单一系统时钟域；复位采用异步置位、同步释放，避免各模块在不同边沿退出复位。",
    55: "第四阶段：编写 soc_firmware.s，加入十项开机自检、shell、数字识别、浮点计算器、Pong 和 Paint，并建立 17 项回归测试；",
    56: "第五阶段：在 ISE 14.7 中综合布局布线并进行真板联调。R16 的 17 项仿真回归、50 MHz 时序约束和主要板端功能均已通过。",
    73: "实现 8 行直接映射和 2 路组相联两种 I-Cache。默认板端采用 2 路组相联结构并使用 LRU 替换；当前冲突地址流仿真中，直接映射 hit=0、miss=12，2 路组相联 hit=10、miss=2，命中率由 0% 提升至 83.3%。",
    84: "为评估功能正确性和架构优化效果，项目使用 Icarus Verilog 建立自动化回归，并由 scripts/analyze.py 统一编译、运行和汇总。R16 测试集包含 17 项 Testbench，覆盖基础指令、流水线冒险、分支预测、RV32M、Custom Float32、自定义位操作、中断、SDRAM、Cache、两层浮点 MLP、表达式计算器、shell 和 SoC I/O，回归率为 100%。",
    101: "端到端数字识别程序（tb_cnn_ablation.v）：评测窗口从串口输入 cnn 命令开始，到 CPU 输出 pred 7 为止。关闭预测时 cycle=70347、CPI=3.021、flush=6960、bp_miss=6740、准确率=6.27%；开启 2-bit BHT 后 cycle=66258、CPI=2.273、flush=432、bp_miss=324、准确率=96.46%。端到端周期降低 5.8%，说明动态预测在真实浮点 MLP 控制流中仍有稳定收益。",
    103: "本项目针对目标器件完成 ISE 14.7 综合、布局布线、静态时序和XPower分析。R16面积、时序与功耗数据均来自2026-07-12的当前实现结果。",
    106: "当前 SoC 系统时钟约束为 20 ns，对应 50 MHz；TEC-PLUS 的 50 MHz 输入经 BUFG 直接供 CPU、Cache、RAM、SDRAM 和外设使用。",
    112: "FPGA 资源利用率（Slice 95% / DSP 0%）",
    114: "R16 使用 Slice Registers 2584（22%）、Slice LUTs 4455（77%）、Occupied Slices 1371（95%）、RAMB16BWER 8（25%）、RAMB8BWER 4（6%）、DSP48A1 0。16 KiB 指令 ROM 正确推断为 4096×32 Block RAM，CPU 寄存器文件保持 Distributed RAM。",
    116: "物理布局布线后的静态时序分析表明，R16 在 50 MHz（20 ns）约束下完成时序收敛。",
    117: "R16 最新 top.twr：Best achievable period = 19.349 ns，Maximum frequency = 51.682 MHz。",
    118: "最坏 setup 裕量为 0.651 ns，Timing errors = 0，Timing Score = 0，所有约束均满足。",
    119: "与 R15 相比，R16 的寄存器和 LUT 数量不变，Occupied Slices 增加 13，RAMB16 增加 4；Fmax 从 54.744 MHz 降至 51.682 MHz。",
    120: "系统仍满足 50 MHz 目标，但 0.651 ns 裕量较小，后续任何组合路径扩展都必须重新运行布局布线和静态时序分析。",
    121: "关键路径已不再由大规模组合乘法器主导。MUL、FMUL 和 FADD 均采用迭代或多周期实现，当前关键限制来自高利用率器件中的存储、前递和控制路径布线。",
    124: "R16 的 20 ns时钟约束通过，Timing errors=0、Score=0；该结果证明扩展模型和计算器后仍可在板卡额定50 MHz下实现。",
    126: "XPower Analyzer基于当前R16的top.ncd/top.pcf完成vector-less活动率传播，处理器系统功耗如下：",
    127: "芯片总供电功耗（Total Supply Power）：91.30 mW。",
    128: "动态功耗（Dynamic Power）：76.19 mW；静态功耗（Static Power）：15.11 mW。",
    129: "XPower置信度为Medium，未加载VCD/SAIF活动文件，因此结果用于PPA版本比较而非板级精密功耗测量。",
    130: "相对R15的90.56 mW，R16总功耗增加0.74 mW（约0.82%），代价主要来自新增指令BRAM和更大的固件功能。",
    132: "图  功耗分解（91.30 mW = 动态76.19 + 静态15.11）",
    134: "R16 已通过迭代算术和 BRAM 映射显著降低 LUT/DSP 压力，同时以 4 个额外 RAMB16 换取 16 KiB 固件空间。",
    135: "当前主要物理风险是 Slice 95% 与 0.651 ns时序裕量，而不是功能完整性；后续优化应优先减少布线拥塞并保持寄存器文件的组合读语义。",
    136: "算术资源复用：RV32M MUL、FMUL32 和 FADD32 使用多周期状态机，DSP48A1 使用量降为0，避免宽组合乘法路径。",
    137: "存储架构：指令 ROM 和片内数据存储映射到 Block RAM；CPU 寄存器文件强制为 Distributed RAM，避免 XST 同步读导致的旧值问题。",
    138: "程序与模型：64→8→10 浮点 MLP 在约2.4 KiB权重空间内提供非线性分类能力；表达式计算器复用 FADD32/FMUL32，并以软件长除法实现 float32 除法。",
    142: "固件 asm/soc_firmware.s 包含十项启动自检，覆盖 RV32I、分支循环、片内 RAM、MUL、DIVU/REMU、FADD32、FMUL32/FGT32、POPCOUNT/BITREVERSE、CSR 和 SDRAM。R16 在 50 MHz 真板上显示 SELFTEST PASS，m0～m9稳定正确，功能验收通过。",
    143: "8×8 数字识别（64→8→10 float32 MLP）",
    144: "用户从 shell 进入 cnn 后，PC 端发送 8×8 二值图像，CPU 执行两层 Custom Float32 MLP：第一层利用二值输入按需累加权重并做 ReLU，第二层完成80次浮点乘加和 argmax。板端精确模型测试集准确率为82.84%，十个演示原型10/10通过。该模型没有卷积层，因此报告中称为MLP数字识别，而不夸大为卷积神经网络。",
    150: "R12 的12.5 MHz历史性能快照只用于证明性能观测链路可用；R16 的架构比较统一采用50 MHz配置下的Icarus计数器和同一测试输入。",
    169: "顶层使用 50 MHz 单一 sys_clk。指令侧为 4096×32 bit（16 KiB）同步片上 ROM，前端默认连接 2 路组相联 I-Cache；数据侧为四个 byte 阵列组成的1024-word片上RAM。MMIO位于0x1000附近，64 MiB外部SDRAM映射到0x10000000～0x13FFFFFF。",
    171: "ISE 14.7 已完成R16综合、布局布线、bitstream和XPower分析。top.twr报告Timing errors=0、Score=0，最小可实现周期19.349 ns，满足50 MHz约束；top.pwr估算总功耗91.30 mW。",
    175: "本项目完成了一套可在 TEC-PLUS Spartan-6 FPGA 上运行的 RV32 五级流水 SoC。R16 固件包含启动自检、串口 shell、两层浮点MLP数字识别、浮点表达式计算器、Pong、Paint、SDRAM和中断演示。最新 Icarus Verilog 回归为17/17 PASS。",
    177: "分支预测微基准中，BHT版本CPI为1.14，baseline为1.57，降低27.4%，预测准确率80%。MLP端到端程序中，baseline为70347 cycles、CPI 3.021；启用BHT后为66258 cycles、CPI 2.273，周期降低5.8%，预测准确率达到96.46%。",
    178: "Cache消融中，冲突地址流的直接映射命中率为0%（hit=0、miss=12），2路组相联加LRU后命中率为83.3%（hit=10、miss=2）。shell和SoC回归同时验证自检结果槽、性能计数器、Pong、Paint、SDRAM、中断和LED状态。",
    180: "R16最新ISE结果为Slice Registers 2584/11440（22%）、Slice LUTs 4455/5720（77%）、Occupied Slices 1371/1430（95%）、RAMB16 8/32（25%）、RAMB8 4/64（6%）、DSP48A1 0。Fmax为51.682 MHz且50 MHz时序通过；XPower总功耗91.30 mW。真板功能全部通过。",
    182: "当前工程风险主要是物理余量：Occupied Slices为95%，50 MHz setup裕量仅0.651 ns。寄存器文件被错误推断为同步BRAM的问题已在R15通过Distributed RAM约束彻底修复，并由R15/R16真板自检验证。",
    183: "实验数据明确区分微基准、端到端程序和系统累计指标；报告中的R16消融均来自同一RTL、同一50 MHz配置和同一输入，避免把人工等待时间混入架构收益。",
    186: "本课程设计从Verilog和RISC-V指令执行出发，在TEC-PLUS Spartan-6 FPGA上完成从CPU数据通路到可交互SoC的实现。系统包含五级流水线、前递与停顿、2-bit BHT、2路I-Cache、片上BRAM、64 MiB SDRAM、UART/LED/KEY、CSR/中断、RV32M、Custom Float32、两层浮点MLP和表达式计算器。",
    188: "通过17项Icarus Verilog自动回归、ISE 14.7实现、静态时序、XPower和50 MHz真板验证，2.0.0版本形成了可复现的RTL—仿真—综合—上板闭环。全部功能通过，R16满足50 MHz约束，模型原型10/10正确，计算器可解析括号和四则运算。",
    189: "后续若继续扩展功能，应优先控制Slice拥塞和关键路径，保持20 ns约束下的正时序裕量；功耗精度可通过导入代表性VCD/SAIF或板级电流测量进一步提高。",
}

for index, text in updates.items():
    replace_paragraph(doc.paragraphs[index], text)

# Cover summary.
summary = ("本课程设计完成 TEC-PLUS FPGA 上的 RV32 五级流水线计算机系统。主要内容包括："
           "RV32I/RV32M、Custom Float32与位操作扩展、数据前递和load-use停顿、2-bit BHT、"
           "2路组相联I-Cache、16 KiB指令ROM、4 KiB数据RAM、64 MiB SDRAM、UART/LED/KEY、"
           "中断和性能计数器；R16固件提供十项开机自检、shell、64→8→10浮点MLP数字识别、"
           "四则表达式计算器、Pong与SDRAM Paint。17项回归及50 MHz真板功能全部通过。")
for row in doc.tables[0].rows:
    for cell in row.cells:
        if "本课程设计完成 TEC-PLUS FPGA" in cell.text:
            replace_cell(cell, summary)

# Regression table: update existing coverage and append the two R16 tests.
regression = doc.tables[1]
for row in regression.rows:
    if row.cells[0].text.strip() == "float":
        replace_cell(row.cells[1], "正负Custom Float32乘加与比较")
existing_tests = {row.cells[0].text.strip() for row in regression.rows}
for name, coverage in (
    ("muldiv-edges", "RV32M符号、高半积、除零与溢出边界"),
    ("calculator", "括号、优先级、小数、负数、除法与错误处理"),
):
    if name in existing_tests:
        continue
    row = regression.add_row()
    replace_cell(row.cells[0], name)
    replace_cell(row.cells[1], coverage)
    replace_cell(row.cells[2], "PASS")

# Ablation table.
ablation = doc.tables[2]
for row in ablation.rows:
    key = row.cells[0].text.replace("\n", " ").strip()
    if "CNN" in key:
        values = [
            "MLP真实程序控制冒险分析", "ENABLE_BP=0\n不开启预测", "ENABLE_BP=1\n开启BHT预测",
            "cycle=70347\nCPI=3.021\nflush=6960\nbp_miss=6740\n准确率=6.27%",
            "cycle=66258\nCPI=2.273\nflush=432\nbp_miss=324\n准确率=96.46%",
            "周期降低5.8%，CPI降低24.8%", "动态预测在两层浮点MLP真实控制流中仍有稳定收益",
        ]
        for cell, value in zip(row.cells, values):
            replace_cell(cell, value)
    elif "I-Cache" in key:
        values = [
            "I-Cache相联度对比", "直接映射（8行）", "2-way组相联+LRU",
            "hit=0, miss=12\n命中率=0%", "hit=10, miss=2\n命中率=83.3%",
            "命中率提升83.3 pct", "2-way结构消除该冲突地址流的反复替换",
        ]
        for cell, value in zip(row.cells, values):
            replace_cell(cell, value)

# Current R16 PPA table.
ppa = doc.tables[3]
ppa_values = [
    ("Slice Registers", "2,584", "11,440", "22%", "流水线、状态机和外设寄存器"),
    ("Slice LUTs", "4,455", "5,720", "77%", "逻辑资源较R13明显下降"),
    ("Occupied Slices", "1,371", "1,430", "95%", "布局成功，仍需控制后续扩展"),
    ("Bonded IOBs", "90", "186", "48%", "SDRAM、LED、KEY、UART与时钟"),
    ("RAMB16BWER", "8", "32", "25%", "16 KiB指令ROM及片上存储"),
    ("RAMB8BWER", "4", "64", "6%", "片上数据存储"),
    ("DSP48A1", "0", "16", "0%", "乘法改为迭代逻辑实现"),
]
while len(ppa.rows) < len(ppa_values) + 1:
    ppa.add_row()
for row, values in zip(ppa.rows[1:], ppa_values):
    for cell, value in zip(row.cells, values):
        replace_cell(cell, value)

doc.save(REPORT)
print(REPORT)
